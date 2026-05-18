"""
Utilities for grounding AI-generated documents in Samiha's real resume files.
"""

from functools import lru_cache
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document


RESUME_ROOT = Path(__file__).parent
# Runtime adds the candidate's own directory via JOB_AGENT_RESUME_DIR env var.
import os as _os
_candidate_dir = Path(_os.getenv("JOB_AGENT_RESUME_DIR", ""))
_RESUME_SEARCH_DIRS = [d for d in [_candidate_dir, RESUME_ROOT, RESUME_ROOT.parent] if d != Path("")]


def _find_resume_file(filename: str) -> Path | None:
    """Search all known locations for a resume file; returns None if not found."""
    if not filename:
        return None
    for base in _RESUME_SEARCH_DIRS:
        candidate = base / filename
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def select_resume_variant(profile: dict, scoring_result: dict) -> tuple[str, Path]:
    """Return the selected resume variant key and file path."""
    variants = profile.get("resume_variants", {})
    # Prefer the scorer's suggestion; fall back to 'default' then 'general_pm'.
    variant_key = scoring_result.get("best_resume_variant", "")
    filename = (
        variants.get(variant_key)
        or variants.get("default")
        or variants.get("general_pm")
        or next(iter(variants.values()), "")
    )
    found = _find_resume_file(filename)
    if found:
        return variant_key or "default", found

    # Try every declared variant in the profile as a fallback.
    for fallback_key, fallback_name in variants.items():
        found = _find_resume_file(fallback_name)
        if found:
            return fallback_key, found

    # Nothing found — return a non-existent file path (caller must handle missing file).
    return variant_key or "default", RESUME_ROOT / (filename or "resume.docx")


@lru_cache(maxsize=16)
def extract_docx_text(path_str: str, char_limit: int = 7000) -> str:
    """Extract concise plain text from a .docx resume."""
    path = Path(path_str)
    if not path.exists() or not path.is_file():
        return ""

    doc = Document(path)
    blocks: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    text = "\n".join(blocks)
    if not text.strip():
        text = _extract_docx_xml_text(path)
    return text[:char_limit]


def _extract_docx_xml_text(path: Path) -> str:
    """Fallback for resumes whose text is stored outside normal paragraphs."""
    try:
        with ZipFile(path) as zf:
            xml_names = [
                name for name in zf.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]
            chunks: list[str] = []
            for name in xml_names:
                root = ET.fromstring(zf.read(name))
                for node in root.iter():
                    if node.tag.endswith('}t') and node.text:
                        chunks.append(node.text.strip())
            return "\n".join(chunk for chunk in chunks if chunk)
    except Exception:
        return ""


def get_selected_resume_source(profile: dict, scoring_result: dict, char_limit: int = 7000) -> dict:
    """Return selected resume metadata plus extracted text for AI prompts."""
    variant_key, path = select_resume_variant(profile, scoring_result)
    return {
        "variant_key": variant_key,
        "filename": path.name,
        "path": str(path),
        "text": extract_docx_text(str(path), char_limit),
    }
