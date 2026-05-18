"""
Document generator — produces tailored .docx resume and cover letter files.
Selects the best matching base template from Samiha's existing resume variants,
then overlays the AI-generated tailored content.
"""

import shutil
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console

console = Console()

# Location of Samiha's existing resume files (parent of linkedin_job_agent/)
RESUME_DIR = Path(__file__).parent.parent.parent


def _safe_filename(text: str, max_len: int = 40) -> str:
    """Convert text to a safe filename segment."""
    safe = "".join(c if c.isalnum() or c in "._- " else "_" for c in text)
    return safe.strip()[:max_len].replace(" ", "_")


def _select_base_resume(profile: dict, scoring_result: dict) -> Path:
    """Pick the best base .docx file for this job based on variant key."""
    variant_key = scoring_result.get("best_resume_variant", "general_pm")
    variants = profile.get("resume_variants", {})
    filename = variants.get(variant_key) or variants.get("general_pm", "")

    candidate_path = RESUME_DIR / filename
    if candidate_path.exists():
        return candidate_path

    # Fallback: pick any existing resume
    for fname in variants.values():
        p = RESUME_DIR / fname
        if p.exists():
            return p

    raise FileNotFoundError(
        f"No resume template found. Expected files in: {RESUME_DIR}\n"
        f"Looked for variant: {variant_key} ({filename})"
    )


def _replace_paragraph_text(para, new_text: str, preserve_style: bool = True):
    """Replace all runs in a paragraph with new text, preserving the first run's format."""
    if not para.runs:
        para.add_run(new_text)
        return
    first_run = para.runs[0]
    font_size = first_run.font.size
    bold = first_run.bold
    italic = first_run.italic
    try:
        color = first_run.font.color.rgb if first_run.font.color.type else None
    except Exception:
        color = None

    for run in para.runs:
        run.text = ""
    first_run.text = new_text
    if preserve_style and font_size:
        first_run.font.size = font_size
    if preserve_style:
        first_run.bold = bold
        first_run.italic = italic
        if color:
            try:
                first_run.font.color.rgb = color
            except Exception:
                pass


def generate_resume(
    profile: dict,
    job_title: str,
    company: str,
    job_id: str,
    scoring_result: dict,
    resume_content: dict,
    output_dir: Path,
) -> Path:
    """
    Generate a tailored resume .docx file.
    Returns path to the generated file.
    """
    output_dir.mkdir(parents=True, exist_ok=True)

    base_path = _select_base_resume(profile, scoring_result)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    company_safe = _safe_filename(company, 25)
    title_safe = _safe_filename(job_title, 30)
    out_name = f"Samiha_Chowdhury_{title_safe}_{company_safe}_{timestamp}.docx"
    out_path = output_dir / out_name

    shutil.copy2(base_path, out_path)
    doc = Document(out_path)

    tailored_title = scoring_result.get("recommended_title_for_resume", job_title)
    summary = resume_content.get("summary", "")
    core_skills = resume_content.get("core_skills", [])
    keyword_tags = resume_content.get("keyword_tags", [])
    key_achievements = resume_content.get("key_achievements", [])
    exp_bullets = resume_content.get("experience_bullets", {})

    # Collect all job bullet lists in order
    exp_bullet_lists = list(exp_bullets.values())

    summary_replaced = False
    skills_replaced = False
    title_replaced = False
    exp_job_idx = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        # Replace document title / headline
        if not title_replaced and any(
            kw in text for kw in ["Project Manager", "Program Manager", "Operations Manager",
                                   "Delivery Manager", "PMO", "Agile"]
        ) and len(text) < 80 and para.runs and (para.runs[0].bold or para.style.name.startswith("Heading")):
            _replace_paragraph_text(para, tailored_title)
            title_replaced = True
            continue

        # Replace summary
        if not summary_replaced and summary and len(text) > 50:
            if any(kw in text.lower() for kw in [
                "experienced", "results-driven", "proven", "accomplished", "dynamic",
                "passionate", "dedicated", "skilled", "professional with"
            ]):
                _replace_paragraph_text(para, summary)
                summary_replaced = True
                continue

        # Replace skills section
        if not skills_replaced and core_skills and "•" in text and len(text) > 40:
            if any(kw in text.lower() for kw in ["agile", "scrum", "jira", "project", "pmp"]):
                skill_text = " • ".join(core_skills[:16])
                _replace_paragraph_text(para, skill_text)
                skills_replaced = True
                continue

    # Overlay experience bullets into the document
    # Find bullet paragraphs (starting with • or -)
    bullet_paras = [
        p for p in doc.paragraphs
        if p.text.strip().startswith("•") or p.text.strip().startswith("-")
    ]

    for i, para in enumerate(bullet_paras):
        text = para.text.strip()
        # Group bullets by job (heuristic: 4 bullets per job)
        job_idx = i // 5
        bullet_idx = i % 5
        if job_idx < len(exp_bullet_lists):
            bullets = exp_bullet_lists[job_idx]
            if bullet_idx < len(bullets):
                new_text = "• " + bullets[bullet_idx]
                _replace_paragraph_text(para, new_text)

    doc.save(out_path)
    console.print(f"  [green]Resume saved:[/green] {out_path.name}")
    return out_path


def generate_cover_letter(
    cover_letter_text: str,
    profile: dict,
    job_title: str,
    company: str,
    job_id: str,
    output_dir: Path,
) -> Path:
    """Generate a cover letter .docx file from the AI-generated text."""
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    company_safe = _safe_filename(company, 25)
    title_safe = _safe_filename(job_title, 30)
    out_name = f"CoverLetter_Samiha_Chowdhury_{title_safe}_{company_safe}_{timestamp}.docx"
    out_path = output_dir / out_name

    personal = profile["personal"]
    doc = Document()

    # Header style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Candidate name heading
    name_para = doc.add_paragraph()
    name_run = name_para.add_run(personal["name"])
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact line
    contact_items = [personal.get("email", ""), personal.get("phone", ""), personal.get("location", "")]
    contact_str = " | ".join(i for i in contact_items if i)
    if contact_str:
        contact_para = doc.add_paragraph(contact_str)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.runs[0].font.size = Pt(10)

    doc.add_paragraph()  # spacer

    # Date
    date_para = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    date_para.runs[0].font.size = Pt(11)

    # Company and role
    doc.add_paragraph()
    role_para = doc.add_paragraph(f"{job_title}\n{company}")
    for run in role_para.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()

    # Cover letter body — split on double newline for paragraphs
    paragraphs = [p.strip() for p in cover_letter_text.split("\n\n") if p.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph(para_text)
        p.paragraph_format.space_after = Pt(12)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save(out_path)
    console.print(f"  [green]Cover letter saved:[/green] {out_path.name}")
    return out_path


def generate_apply_pack(
    profile: dict,
    job_title: str,
    company: str,
    job_location: str,
    job_id: str,
    scoring_result: dict,
    resume_content: dict,
    cover_letter_text: str,
    config: dict,
) -> tuple[Path, Path]:
    """
    Generate both resume and cover letter for a job.
    Returns (resume_path, cover_letter_path).
    """
    resume_dir = Path(config["output"]["resumes_dir"])
    cl_dir = Path(config["output"]["cover_letters_dir"])

    resume_path = generate_resume(
        profile=profile,
        job_title=job_title,
        company=company,
        job_id=job_id,
        scoring_result=scoring_result,
        resume_content=resume_content,
        output_dir=resume_dir,
    )

    cl_path = generate_cover_letter(
        cover_letter_text=cover_letter_text,
        profile=profile,
        job_title=job_title,
        company=company,
        job_id=job_id,
        output_dir=cl_dir,
    )

    return resume_path, cl_path
